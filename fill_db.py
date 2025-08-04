import chromadb
import os
import glob
from langchain_text_splitters import RecursiveCharacterTextSplitter
import re
import json

# setting the environment
CHROMA_PATH = r"chroma_db"
COLLECTION_NAME = "universal_content"

# Supported file types and their extensions
SUPPORTED_EXTENSIONS = {
    'text': ['.txt', '.md', '.py', '.js', '.html', '.css', '.sql', '.sh', '.bat', '.yml', '.yaml', '.xml'],
    'pdf': ['.pdf'],
    'json': ['.json'],
    'csv': ['.csv'],
    'word': ['.docx'],
    'data': ['.log', '.ini', '.cfg', '.conf']
}

# Enhanced content analysis patterns
PATTERNS = {
    'procedures': [
        r'step \d+', r'step\s*\d+:', r'\d+\.\s+', r'first|second|third|next|then|finally',
        r'procedure', r'instructions', r'how to', r'guide', r'getting started',
        r'tutorial', r'walkthrough', r'setup', r'installation'
    ],
    'lists': [
        r'^\s*[-*•]\s+', r'^\s*\d+\.\s+', r'types?\s+of', r'includes?:',
        r'following:', r'such as:', r'examples?:', r'supports.*types?',
        r'two types', r'three types', r'several types', r'types\s+of\s+workflows',
        r'categories', r'options', r'methods', r'approaches'
    ],
    'technical': [
        r'api', r'function', r'class', r'method', r'parameter', r'return',
        r'configuration', r'settings', r'system', r'technical', r'def\s+\w+',
        r'import\s+', r'from\s+\w+', r'#!/', r'<html>', r'SELECT\s+', r'CREATE\s+'
    ],
    'faq': [
        r'q:', r'question:', r'a:', r'answer:', r'what\s+is', r'how\s+do',
        r'why\s+does', r'can\s+i', r'frequently', r'common', r'\?\s*$'
    ],
    'conversational': [
        r'\bi\s+', r'\byou\s+', r'\bwe\s+', r'let\'s', r'here\'s',
        r'first\s+', r'now\s+', r'okay', r'welcome', r'hello'
    ],
    'code': [
        r'def\s+\w+', r'class\s+\w+', r'function\s+\w+', r'var\s+\w+',
        r'const\s+\w+', r'let\s+\w+', r'import\s+', r'from\s+\w+\s+import',
        r'#!/usr/bin', r'<\?php', r'<!DOCTYPE', r'public\s+class'
    ],
    'data': [
        r'^\s*\w+:', r'^\s*"\w+":', r'^\w+,\w+', r'\{.*\}', r'\[.*\]',
        r'=\s*\w+', r':\s*\w+', r'csv', r'json', r'database'
    ]
}

def read_file_content(file_path):
    """Universal file reader that handles different file types"""
    
    filename = os.path.basename(file_path)
    extension = os.path.splitext(filename)[1].lower()
    
    print(f"📖 Reading: {filename} (type: {extension})")
    
    try:
        # Text-based files (most common)
        if extension in SUPPORTED_EXTENSIONS['text']:
            try:
                with open(file_path, 'r', encoding='utf-8') as f:
                    content = f.read().strip()
                return content, 'text'
            except UnicodeDecodeError:
                with open(file_path, 'r', encoding='latin-1') as f:
                    content = f.read().strip()
                print(f"   Used latin-1 encoding")
                return content, 'text'
        
        # PDF files
        elif extension in SUPPORTED_EXTENSIONS['pdf']:
            try:
                from PyPDF2 import PdfReader
                with open(file_path, 'rb') as f:
                    reader = PdfReader(f)
                    content = ""
                    for page in reader.pages:
                        content += page.extract_text() + "\n"
                print(f"   Extracted {len(reader.pages)} pages")
                return content.strip(), 'pdf'
            except ImportError:
                print(f"   ⚠️  PyPDF2 not installed, skipping PDF")
                return None, None
            except Exception as e:
                print(f"   ❌ PDF reading error: {str(e)}")
                return None, None
        
        # JSON files
        elif extension in SUPPORTED_EXTENSIONS['json']:
            with open(file_path, 'r', encoding='utf-8') as f:
                data = json.load(f)
            # Convert JSON to readable text
            content = json.dumps(data, indent=2)
            # Also add a summary at the top
            summary = f"JSON Data Structure:\n"
            if isinstance(data, dict):
                summary += f"Object with {len(data)} keys: {', '.join(list(data.keys())[:5])}\n\n"
            elif isinstance(data, list):
                summary += f"Array with {len(data)} items\n\n"
            content = summary + content
            return content, 'json'
        
        # CSV files
        elif extension in SUPPORTED_EXTENSIONS['csv']:
            try:
                import csv
                content = ""
                with open(file_path, 'r', encoding='utf-8') as f:
                    reader = csv.reader(f)
                    headers = next(reader, None)
                    if headers:
                        content += f"CSV Data with columns: {', '.join(headers)}\n\n"
                        content += ','.join(headers) + '\n'
                        # Read first 10 rows as sample
                        for i, row in enumerate(reader):
                            if i >= 10:
                                content += f"... ({i} more rows)\n"
                                break
                            content += ','.join(row) + '\n'
                return content, 'csv'
            except Exception as e:
                print(f"   ❌ CSV reading error: {str(e)}")
                return None, None
        
        # Word documents
        elif extension in SUPPORTED_EXTENSIONS['word']:
            try:
                from docx import Document
                doc = Document(file_path)
                content = ""
                for paragraph in doc.paragraphs:
                    content += paragraph.text + "\n"
                print(f"   Extracted {len(doc.paragraphs)} paragraphs")
                return content.strip(), 'word'
            except ImportError:
                print(f"   ⚠️  python-docx not installed, skipping Word document")
                return None, None
            except Exception as e:
                print(f"   ❌ Word document reading error: {str(e)}")
                return None, None
        
        # Data/config files
        elif extension in SUPPORTED_EXTENSIONS['data']:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read().strip()
            return content, 'data'
        
        else:
            # Try to read as text file anyway
            try:
                with open(file_path, 'r', encoding='utf-8') as f:
                    content = f.read().strip()
                print(f"   ⚠️  Unknown extension, treating as text")
                return content, 'unknown'
            except Exception as e:
                print(f"   ❌ Cannot read file: {str(e)}")
                return None, None
                
    except Exception as e:
        print(f"   ❌ Error reading {filename}: {str(e)}")
        return None, None

def analyze_file_characteristics(file_path, content, file_type):
    """Analyze file size and content to determine optimal chunking strategy"""
    
    filename = os.path.basename(file_path)
    file_size = len(content)
    
    print(f"\n📊 ANALYZING: {filename}")
    print("-" * 50)
    
    # Basic metrics
    words = len(content.split())
    lines = content.count('\n') + 1
    paragraphs = len([p for p in content.split('\n\n') if p.strip()])
    
    print(f"📏 Size: {file_size:,} characters")
    print(f"📝 Words: {words:,}")
    print(f"📋 Lines: {lines}")
    print(f"📄 Paragraphs: {paragraphs}")
    print(f"🗂️  File type: {file_type}")
    
    # Content structure analysis
    avg_line_length = file_size / lines if lines > 0 else 0
    avg_paragraph_length = file_size / paragraphs if paragraphs > 0 else 0
    
    print(f"📊 Avg line length: {avg_line_length:.1f} chars")
    print(f"📊 Avg paragraph length: {avg_paragraph_length:.1f} chars")
    
    # Content type detection
    content_lower = content.lower()
    content_types = []
    
    for content_type, patterns in PATTERNS.items():
        score = 0
        for pattern in patterns:
            try:
                matches = len(re.findall(pattern, content_lower, re.MULTILINE))
                score += matches
            except re.error:
                continue  # Skip invalid regex patterns
        
        # Apply priority weights and file type bonuses
        if content_type in ['procedures', 'lists'] and score > 0:
            score *= 2  # Double weight for procedures and lists
        elif content_type == 'technical' and score > 0:
            score *= 1.5  # Higher weight for technical content
        elif content_type == 'code' and file_type in ['text'] and score > 0:
            score *= 2  # Code files are important
        elif content_type == 'data' and file_type in ['json', 'csv', 'data'] and score > 0:
            score *= 3  # Data files get data type bonus
        
        if score > 0:
            content_types.append((content_type, score))
    
    # Sort by relevance (now with weighted scores)
    content_types.sort(key=lambda x: x[1], reverse=True)
    
    if content_types:
        primary_type = content_types[0][0]
        print(f"🎯 Detected content type: {primary_type}")
        if len(content_types) > 1:
            print(f"   Secondary types: {', '.join([ct[0] for ct in content_types[1:3]])}")
    else:
        primary_type = "general"
        print(f"🎯 Content type: general")
    
    return {
        'filename': filename,
        'file_size': file_size,
        'file_type': file_type,
        'words': words,
        'lines': lines,
        'paragraphs': paragraphs,
        'avg_line_length': avg_line_length,
        'avg_paragraph_length': avg_paragraph_length,
        'primary_type': primary_type,
        'content_types': content_types
    }

def determine_chunking_strategy(analysis):
    """Determine optimal chunking parameters based on file analysis"""
    
    size = analysis['file_size']
    primary_type = analysis['primary_type']
    file_type = analysis['file_type']
    avg_para = analysis['avg_paragraph_length']
    
    print(f"\n🎯 CHUNKING STRATEGY FOR: {analysis['filename']}")
    print("-" * 50)
    
    # Base strategy on file size
    if size < 1000:  # Small files
        base_chunk_size = 300
        strategy = "small_file"
        print("📏 File size: Small (< 1KB)")
        
    elif size < 5000:  # Medium files
        base_chunk_size = 500
        strategy = "medium_file"
        print("📏 File size: Medium (1-5KB)")
        
    elif size < 20000:  # Large files
        base_chunk_size = 800
        strategy = "large_file"
        print("📏 File size: Large (5-20KB)")
        
    elif size < 100000:  # Very large files
        base_chunk_size = 1200
        strategy = "xlarge_file"
        print("📏 File size: Extra Large (20-100KB)")
        
    else:  # Huge files
        base_chunk_size = 1600
        strategy = "huge_file"
        print("📏 File size: Huge (> 100KB)")
    
    # Adjust based on file type
    file_type_adjustments = {
        'pdf': 1.3,     # PDFs often need larger chunks
        'code': 1.2,    # Code needs context
        'json': 0.8,    # JSON can be chunked smaller
        'csv': 0.9,     # CSV data can be smaller
        'word': 1.2,    # Word docs often have complex formatting
        'data': 1.0     # Config files standard size
    }
    
    if file_type in file_type_adjustments:
        multiplier = file_type_adjustments[file_type]
        base_chunk_size = int(base_chunk_size * multiplier)
        print(f"🗂️  File type adjustment: {file_type} (x{multiplier})")
    
    # Adjust based on content type
    type_adjustments = {
        'procedures': 1.5,    # Need larger chunks for step-by-step
        'lists': 1.4,        # Keep list items together
        'technical': 1.3,    # Technical context important
        'code': 1.4,         # Code needs context
        'data': 1.1,         # Data structures need some context
        'faq': 0.8,          # Q&A can be smaller
        'conversational': 0.7  # Chat-like content
    }
    
    if primary_type in type_adjustments:
        multiplier = type_adjustments[primary_type]
        base_chunk_size = int(base_chunk_size * multiplier)
        print(f"🔧 Content type adjustment: {primary_type} (x{multiplier})")
    
    # Adjust based on paragraph structure
    if avg_para > 0:
        if avg_para < 150:  # Very short paragraphs
            base_chunk_size = max(base_chunk_size, 400)  # Ensure minimum size
            print("📄 Short paragraphs detected → Minimum chunk size enforced")
        elif avg_para > 800:  # Very long paragraphs
            base_chunk_size = max(base_chunk_size, 1000)  # Ensure large enough
            print("📄 Long paragraphs detected → Larger chunks needed")
    
    # Calculate overlap (20-25% of chunk size)
    overlap = int(base_chunk_size * 0.25)
    
    # Determine number of results based on expected chunks
    expected_chunks = max(2, size // base_chunk_size)
    if expected_chunks <= 3:
        n_results = 3
    elif expected_chunks <= 8:
        n_results = min(6, expected_chunks)
    else:
        n_results = 8
    
    print(f"📏 Chunk size: {base_chunk_size}")
    print(f"🔄 Overlap: {overlap}")
    print(f"🔍 Recommended n_results: {n_results}")
    print(f"📊 Expected chunks: ~{expected_chunks}")
    
    return {
        'chunk_size': base_chunk_size,
        'overlap': overlap,
        'n_results': n_results,
        'strategy': strategy,
        'expected_chunks': expected_chunks
    }

def adaptive_chunk_content(content, strategy_params, metadata):
    """Chunk content using adaptive parameters"""
    
    chunk_size = strategy_params['chunk_size']
    overlap = strategy_params['overlap']
    
    # Create text splitter with adaptive parameters
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=overlap,
    length_function=len,
    is_separator_regex=False,
)

    # Create document object
    class Document:
        def __init__(self, page_content, metadata):
            self.page_content = page_content
            self.metadata = metadata
    
    doc = Document(content, metadata)
    chunks = splitter.split_documents([doc])
    
    return chunks

def find_all_supported_files(directory):
    """Find all supported files in a directory"""
    
    if not os.path.exists(directory):
        return []
    
    all_files = []
    all_extensions = []
    for ext_list in SUPPORTED_EXTENSIONS.values():
        all_extensions.extend(ext_list)
    
    for ext in all_extensions:
        pattern = os.path.join(directory, f"*{ext}")
        files = glob.glob(pattern)
        all_files.extend(files)
    
    return sorted(all_files)

def process_directory(directory_name, start_chunk_id=0):
    """Process all supported files in a directory"""
    
    files = find_all_supported_files(directory_name)
    
    if not files:
        return None
    
    print(f"📁 Found {len(files)} supported files:")
    file_types = {}
    for file_path in files:
        filename = os.path.basename(file_path)
        extension = os.path.splitext(filename)[1].lower()
        file_size = os.path.getsize(file_path)
        
        file_types[extension] = file_types.get(extension, 0) + 1
        print(f"  📄 {filename} ({file_size} bytes, {extension})")
    
    print(f"\n📊 File type summary:")
    for ext, count in sorted(file_types.items()):
        print(f"  {ext}: {count} files")
    
    # Process each file
    all_documents = []
    all_metadata = []
    all_ids = []
    chunk_id = start_chunk_id
    
    print(f"\n🔍 UNIVERSAL CONTENT ANALYSIS")
    print("=" * 60)
    
    for file_path in files:
        filename = os.path.basename(file_path)
        
        try:
            # Read file content
            content, file_type = read_file_content(file_path)
            
            if not content:
                print(f"⚠️  Skipping {filename} - no content extracted")
                continue
            
            # Analyze content
            analysis = analyze_file_characteristics(file_path, content, file_type)
            
            # Determine chunking strategy
            strategy = determine_chunking_strategy(analysis)
            
            # Extract title
            lines = content.split('\n')
            title = lines[0].strip()[:100] if lines else filename
            
            # Create metadata
            base_metadata = {
                'source': filename,
                'title': title,
                'file_path': file_path,
                'file_size': analysis['file_size'],
                'file_type': file_type,
                'content_type': analysis['primary_type'],
                'chunking_strategy': strategy['strategy']
            }
            
            # Chunk content
            chunks = adaptive_chunk_content(content, strategy, base_metadata)
            
            print(f"✂️  Created {len(chunks)} chunks (strategy: {strategy['strategy']})")
            
            # Add to collections
            for i, chunk in enumerate(chunks):
                all_documents.append(chunk.page_content)
                # Create unique ID using filename and chunk index
                clean_filename = filename.replace('.', '_').replace(' ', '_')
                unique_id = f"UNIVERSAL_{clean_filename}_{i}_{chunk_id}"
                all_ids.append(unique_id)
                
                chunk_metadata = base_metadata.copy()
                chunk_metadata.update({
                    'chunk_id': chunk_id,
                    'chunk_index': i,
                    'unique_id': unique_id,
                    'actual_chunk_size': len(chunk.page_content)
                })
                all_metadata.append(chunk_metadata)
                
                chunk_id += 1
        
        except Exception as e:
            print(f"❌ Error processing {filename}: {str(e)}")
            continue
    
    return all_documents, all_metadata, all_ids

def main():
    """Main universal ingestion function"""
    
    print("🌍 UNIVERSAL CONTENT INGESTION SYSTEM")
    print("=" * 60)
    print("Supports: Text, Markdown, PDF, Word, JSON, CSV, Code files and more!")
    print()
    
    # Show supported file types
    print("🗂️  Supported file types:")
    for category, extensions in SUPPORTED_EXTENSIONS.items():
        ext_str = ', '.join(extensions)
        print(f"   {category.upper()}: {ext_str}")
    print()
    
    # Get directories to process
    print("📁 Choose content sources:")
    print("1. Current directory")
    print("2. Specific directories")
    print("3. Default directories (txt_data, md_data, data)")
    
    try:
        choice = int(input("Enter choice (1-3): "))
    except ValueError:
        choice = 3
    
    directories = []
    if choice == 1:
        directories = ["."]
    elif choice == 2:
        print("Enter directory names (one per line, empty line to finish):")
        while True:
            dir_name = input("Directory: ").strip()
            if not dir_name:
                break
            directories.append(dir_name)
    else:  # Default
        directories = ["txt_data", "md_data", "data"]
        # Create directories if they don't exist
        for dir_name in directories:
            if not os.path.exists(dir_name):
                os.makedirs(dir_name)
                print(f"✅ Created directory: {dir_name}")
    
    # Initialize ChromaDB
    chroma_client = chromadb.PersistentClient(path=CHROMA_PATH)
    
    # Ask about replacing data
    replace_data = input(f"\nReplace existing data in collection '{COLLECTION_NAME}'? (y/n): ").strip().lower()
    
    if replace_data in ['y', 'yes']:
        try:
            chroma_client.delete_collection(name=COLLECTION_NAME)
            print("✅ Deleted existing collection")
        except:
            print("No existing collection to delete")
    
    # Create collection
    try:
        collection = chroma_client.create_collection(
            name=COLLECTION_NAME,
            metadata={"hnsw:num_threads": 1}
        )
        print("✅ Created new collection")
    except:
        collection = chroma_client.get_collection(name=COLLECTION_NAME)
        print("✅ Using existing collection")
    
    # Process all directories
    all_docs = []
    all_meta = []
    all_ids = []
    global_chunk_id = 0
    
    for directory in directories:
        if not os.path.exists(directory):
            print(f"⚠️  Directory {directory} does not exist, skipping")
            continue
            
        print(f"\n📂 PROCESSING DIRECTORY: {directory}")
        print("=" * 50)
        
        result = process_directory(directory, global_chunk_id)
        if result:
            docs, meta, ids = result
            all_docs.extend(docs)
            all_meta.extend(meta)
            all_ids.extend(ids)
            global_chunk_id += len(docs)
            print(f"✅ Added {len(docs)} chunks from {directory}")
    
    if not all_docs:
        print("❌ No content was processed!")
        print("Make sure you have supported files in your directories.")
        return
    
    # Add to ChromaDB
    print(f"\n💾 ADDING CONTENT TO DATABASE")
    print("=" * 40)

    collection.upsert(
        documents=all_docs,
        metadatas=all_meta,
        ids=all_ids
    )
    
    print(f"✅ Successfully added {len(all_docs)} chunks to the database!")
    
    # Show comprehensive summary
    strategy_summary = {}
    file_type_summary = {}
    content_type_summary = {}
    
    for meta in all_meta:
        strategy = meta.get('chunking_strategy', 'unknown')
        file_type = meta.get('file_type', 'unknown')
        content_type = meta.get('content_type', 'unknown')
        
        strategy_summary[strategy] = strategy_summary.get(strategy, 0) + 1
        file_type_summary[file_type] = file_type_summary.get(file_type, 0) + 1
        content_type_summary[content_type] = content_type_summary.get(content_type, 0) + 1
    
    print(f"\n📊 PROCESSING SUMMARY:")
    print(f"   📄 Total files processed: {len(set(meta.get('source', '') for meta in all_meta))}")
    print(f"   📝 Total chunks created: {len(all_docs)}")
    
    print(f"\n🗂️  File types processed:")
    for ftype, count in sorted(file_type_summary.items(), key=lambda x: x[1], reverse=True):
        print(f"   {ftype}: {count} chunks")
    
    print(f"\n🎯 Content types detected:")
    for ctype, count in sorted(content_type_summary.items(), key=lambda x: x[1], reverse=True):
        print(f"   {ctype}: {count} chunks")
    
    print(f"\n⚙️  Chunking strategies used:")
    for strategy, count in sorted(strategy_summary.items(), key=lambda x: x[1], reverse=True):
        print(f"   {strategy}: {count} chunks")
    
    print(f"\n🚀 Universal ingestion complete!")
    print(f"📁 Collection name: {COLLECTION_NAME}")
    print("Run 'python ask.py' to query your intelligently processed content!")

if __name__ == "__main__":
    main()